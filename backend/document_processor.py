"""
Document Processor Module
Handles extraction of text and data from various file formats
"""

import os
from typing import Dict, Optional
from pathlib import Path

# Import libraries with error handling
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


class DocumentProcessor:
    """Process different document types and extract text"""
    
    def __init__(self):
        """Initialize the document processor"""
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.csv': self._process_csv,
            '.docx': self._process_word,
            '.txt': self._process_text
        }
    
    def process(self, file_path: str) -> Dict:
        """
        Process a document and extract text
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary with extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = file_path.suffix.lower()
        
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported file type: {ext}. Supported types: {', '.join(self.supported_formats.keys())}")
        
        processor_func = self.supported_formats[ext]
        return processor_func(file_path)
    
    def _process_pdf(self, file_path: Path) -> Dict:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        try:
            text = ""
            page_count = 0
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                page_count = len(reader.pages)
                
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    text += f"\n\n[Page {i + 1}]\n{page_text}"
            
            return {
                'type': 'pdf',
                'text': text.strip(),
                'pages': page_count,
                'success': True
            }
        except Exception as e:
            raise Exception(f"Failed to process PDF: {str(e)}")
    
    def _process_excel(self, file_path: Path) -> Dict:
        """Extract data from Excel file"""
        if not PANDAS_AVAILABLE:
            raise ImportError("pandas and openpyxl not installed. Install with: pip install pandas openpyxl")
        
        try:
            # Read all sheets
            excel_file = pd.read_excel(file_path, sheet_name=None)
            
            text = ""
            sheet_count = 0
            
            for sheet_name, df in excel_file.items():
                sheet_count += 1
                text += f"\n\n[Sheet: {sheet_name}]\n"
                text += f"Columns: {', '.join(df.columns)}\n\n"
                text += df.to_string(index=False)
            
            return {
                'type': 'excel',
                'text': text.strip(),
                'sheets': sheet_count,
                'pages': sheet_count,
                'success': True
            }
        except Exception as e:
            raise Exception(f"Failed to process Excel file: {str(e)}")
    
    def _process_csv(self, file_path: Path) -> Dict:
        """Extract data from CSV file"""
        if not PANDAS_AVAILABLE:
            raise ImportError("pandas not installed. Install with: pip install pandas")
        
        try:
            df = pd.read_csv(file_path)
            
            text = f"Columns: {', '.join(df.columns)}\n\n"
            text += df.to_string(index=False)
            
            return {
                'type': 'csv',
                'text': text.strip(),
                'rows': len(df),
                'pages': 1,
                'success': True
            }
        except Exception as e:
            raise Exception(f"Failed to process CSV file: {str(e)}")
    
    def _process_word(self, file_path: Path) -> Dict:
        """Extract text from Word document"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            doc = DocxDocument(file_path)
            
            text = ""
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    text += para.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                text += "\n[Table]\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
            
            return {
                'type': 'docx',
                'text': text.strip(),
                'paragraphs': len(doc.paragraphs),
                'pages': max(1, len(doc.paragraphs) // 20),  # Rough estimate
                'success': True
            }
        except Exception as e:
            raise Exception(f"Failed to process Word document: {str(e)}")
    
    def _process_text(self, file_path: Path) -> Dict:
        """Read plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            
            line_count = text.count('\n') + 1
            
            return {
                'type': 'txt',
                'text': text,
                'lines': line_count,
                'pages': max(1, line_count // 50),  # Rough estimate
                'success': True
            }
        except Exception as e:
            raise Exception(f"Failed to process text file: {str(e)}")
    
    def get_supported_formats(self) -> list:
        """Return list of supported file formats"""
        return list(self.supported_formats.keys())

